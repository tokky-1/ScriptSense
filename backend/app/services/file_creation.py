from pathlib import Path


def create_txt(pages: list[str], output_path: Path) -> Path:
    separator = "\n" + "=" * 60 + "\n\n"
    output_path.write_text(separator.join(pages), encoding="utf-8")
    return output_path


def create_pdf(pages: list[str], output_path: Path) -> Path:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=2.5 * cm,
        rightMargin=2.5 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm,
    )
    base = getSampleStyleSheet()["BodyText"]
    body = ParagraphStyle(
        "ScriptSenseBody",
        parent=base,
        fontName="Helvetica",
        fontSize=11,
        leading=16,
        spaceAfter=4,
    )

    story = []
    for i, page_text in enumerate(pages):
        if i > 0:
            story.append(PageBreak())
        for line in page_text.split("\n"):
            story.append(Paragraph(line.strip() or "&nbsp;", body))

    doc.build(story)
    return output_path


def create_docx(pages: list[str], output_path: Path) -> Path:
    from docx import Document
    from docx.enum.text import WD_BREAK
    from docx.shared import Cm, Pt

    document = Document()
    for section in document.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    for i, page_text in enumerate(pages):
        if i > 0:
            document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        for line in page_text.split("\n"):
            p = document.add_paragraph(line)
            for run in p.runs:
                run.font.size = Pt(11)

    document.save(str(output_path))
    return output_path


def create_file(pages: list[str], output_format: str, output_path: Path) -> Path:
    fmt = output_format.lower().lstrip(".")
    if fmt == "txt":
        return create_txt(pages, output_path)
    if fmt == "pdf":
        return create_pdf(pages, output_path)
    if fmt in ("doc", "docx"):
        return create_docx(pages, output_path)
    raise ValueError(f"Unsupported output format: {output_format!r}")
